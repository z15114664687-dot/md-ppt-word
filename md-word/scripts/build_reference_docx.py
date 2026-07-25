#!/usr/bin/env python3
"""Build the A4 reference DOCX used by the md-word pipeline.

All layout values come from shared-assets/design-tokens.json; this script must
not hardcode margins, colors, or font sizes.
"""

from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TOKENS_PATH = ROOT / "shared-assets" / "design-tokens.json"


def load_word_tokens() -> dict:
    return json.loads(TOKENS_PATH.read_text(encoding="utf-8"))["word"]


def pandoc_command() -> list[str]:
    if shutil.which("pandoc"):
        return [shutil.which("pandoc") or "pandoc"]
    if shutil.which("quarto"):
        # Quarto 自带一份 Pandoc；Word 线本身不需要 Quarto。
        return [shutil.which("quarto") or "quarto", "pandoc"]
    raise RuntimeError("缺少 Pandoc（或 Quarto 内置 Pandoc）：当前只能预检 Markdown，不能生成 DOCX。安装方式见仓库 README.md")


def new_document_from_pandoc_defaults() -> Document:
    """Open Pandoc's built-in reference.docx as the base document.

    Pandoc 的 docx writer 会引用 Compact/Table/First Paragraph 等自带样式；
    参考模板必须包含它们，否则 Word/LibreOffice 对缺失样式的回退会破坏排版。
    """
    data = subprocess.run(
        [*pandoc_command(), "--print-default-data-file", "reference.docx"],
        check=True,
        capture_output=True,
    ).stdout
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        handle.write(data)
        base_path = Path(handle.name)
    document = Document(base_path)
    base_path.unlink()
    return document


def _set_font(style, east_asia: str, latin: str, size: float, bold: bool = False, color: str | None = None) -> None:
    # LibreOffice on macOS does not consistently honor the eastAsia override
    # when ascii/hAnsi point to a Latin-only font, so the preview profile uses
    # the CJK family (which bundles Latin glyphs) for every script.
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    rfonts = style._element.rPr.rFonts
    for attribute in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attribute), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    # Pandoc 模板的样式带 theme 字体属性（asciiTheme 等），优先级高于显式
    # 字体名，必须删除，否则标题仍按主题字体（Cambria 系）渲染。
    for attribute in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rfonts.get(qn(attribute)) is not None:
            del rfonts.attrib[qn(attribute)]
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def _style(document: Document, name: str, kind=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, kind)


def _set_first_line_indent_chars(style, chars: int, body_size_pt: float) -> None:
    """Set a CJK char-based first-line indent with a twips fallback."""
    p_pr = style._element.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:firstLine"), str(int(chars * body_size_pt * 20)))
    ind.set(qn("w:firstLineChars"), str(chars * 100))


def _clear_first_line_indent(style) -> None:
    p_pr = style._element.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:firstLineChars"), "0")


def field(paragraph, instruction: str, cached_lines: list[str] | None = None) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instruction_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instruction_run._r.append(instr)
    separator_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separator_run._r.append(separate)
    for index, line in enumerate(cached_lines or []):
        cached_run = paragraph.add_run(line)
        if index + 1 < len(cached_lines or []):
            cached_run.add_break()
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def configure_document(document: Document, font_profile: str = "preview") -> None:
    tokens = load_word_tokens()
    if font_profile not in tokens["font_profiles"]:
        raise ValueError(f"未知字体档位：{font_profile}，可选 {sorted(tokens['font_profiles'])}")
    fonts = tokens["font_profiles"][font_profile]
    sizes = tokens["font_sizes_pt"]
    colors = tokens["colors"]
    margins = tokens["margins_mm"]
    primary = colors["primary"].lstrip("#")
    body_zh, body_latin = fonts["body_zh"], fonts["body_latin"]
    head_zh, head_latin = fonts["heading_zh"], fonts["heading_latin"]

    for section in document.sections:
        section.page_width = Mm(tokens["page"]["width_mm"])
        section.page_height = Mm(tokens["page"]["height_mm"])
        section.top_margin = Mm(margins["top"])
        section.right_margin = Mm(margins["right"])
        section.bottom_margin = Mm(margins["bottom"])
        section.left_margin = Mm(margins["left"])
        section.header_distance = Mm(tokens["header_distance_mm"])
        section.footer_distance = Mm(tokens["footer_distance_mm"])
        for paragraph in section.header.paragraphs:
            paragraph.text = ""
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        field(paragraph, "PAGE", ["1"])

    normal = document.styles["Normal"]
    _set_font(normal, body_zh, body_latin, sizes["body"])
    normal.paragraph_format.line_spacing = Pt(tokens["line_spacing_pt"])
    normal.paragraph_format.space_after = Pt(tokens["space_after_pt"])

    # Pandoc styles ordinary paragraphs as Body Text / First Paragraph; the
    # CJK first-line indent lives there so table cells (Compact) stay flush.
    body_text = _style(document, "Body Text")
    body_text.base_style = normal
    first_paragraph = _style(document, "First Paragraph")
    first_paragraph.base_style = body_text
    indent_chars = tokens["body_first_line_indent_chars"]
    for style in (body_text, first_paragraph):
        _set_first_line_indent_chars(style, indent_chars, sizes["body"])
    compact = _style(document, "Compact")
    compact.base_style = body_text
    _clear_first_line_indent(compact)

    title = document.styles["Title"]
    _set_font(title, head_zh, head_latin, sizes["cover_title"], True, primary)
    title_spacing = tokens["style_spacing_pt"]["Title"]
    title.paragraph_format.space_before = Pt(title_spacing["before"])
    title.paragraph_format.space_after = Pt(title_spacing["after"])
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size in (
        ("Heading 1", sizes["heading_1"]),
        ("Heading 2", sizes["heading_2"]),
        ("Heading 3", sizes["body"]),
    ):
        style = document.styles[name]
        _set_font(style, head_zh, head_latin, size, True, primary)
        spacing = tokens["style_spacing_pt"][name]
        style.paragraph_format.space_before = Pt(spacing["before"])
        style.paragraph_format.space_after = Pt(spacing["after"])
        style.paragraph_format.keep_with_next = True
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True

    # 国金参考样式：图和表共用"图表N："编号，题注统一左对齐置于对象上方。
    for name, size, bold in (
        ("Exhibit Title", sizes["caption"], True),
        ("Source Note", sizes["source_note"], False),
        ("Appendix Heading", sizes["heading_1"], True),
        ("Directory Heading", sizes["heading_1"], True),
    ):
        style = _style(document, name)
        _set_font(
            style,
            body_zh if name == "Source Note" else head_zh,
            body_latin if name == "Source Note" else head_latin,
            size,
            bold,
            primary if "Heading" in name else None,
        )
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        spacing = tokens["style_spacing_pt"][name]
        style.paragraph_format.space_before = Pt(spacing["before"])
        style.paragraph_format.space_after = Pt(spacing["after"])
        style.paragraph_format.keep_with_next = name != "Source Note"
        if name == "Appendix Heading":
            style.paragraph_format.page_break_before = True
            # 大纲级别 0：附录要进内容目录（TOC \o 按大纲级别收集）。
            p_pr = style._element.get_or_add_pPr()
            outline = p_pr.find(qn("w:outlineLvl"))
            if outline is None:
                outline = OxmlElement("w:outlineLvl")
                p_pr.append(outline)
            outline.set(qn("w:val"), "0")
        _clear_first_line_indent(style)

    settings = document.settings._element
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def build(output: Path, font_profile: str = "preview") -> None:
    tokens = load_word_tokens()
    header_fill = tokens["colors"]["table_header"].lstrip("#")
    header_text = tokens["colors"]["table_header_text"].lstrip("#")
    document = new_document_from_pandoc_defaults()
    configure_document(document, font_profile)
    document.add_paragraph("研究报告参考样式", style="Title")
    document.add_paragraph("内容目录", style="Directory Heading")
    field(document.add_paragraph(style="Normal"), r'TOC \o "1-3" \h \z \u', ["一、研究结论"])
    document.add_paragraph("图表目录", style="Directory Heading")
    field(
        document.add_paragraph(style="Normal"),
        r'TOC \h \z \t "Exhibit Title,1"',
        ["图表1：示例图题", "图表2：示例表题"],
    )
    document.add_paragraph("一、研究结论", style="Heading 1")
    document.add_paragraph("本页仅用于检查字体、间距、题注与表格样式；Pandoc 使用 reference.docx 时不会复制这些示例内容。", style="Body Text")
    document.add_paragraph("图表1：示例图题", style="Exhibit Title")
    document.add_paragraph("来源：模拟数据，仅作演示。", style="Source Note")
    document.add_paragraph("图表2：示例表题", style="Exhibit Title")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "数值"
    table.cell(1, 0).text = "年化收益"
    table.cell(1, 1).text = "10.2%"
    header = table.rows[0]
    tr_pr = header._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for cell in header.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), header_fill)
        tc_pr.append(shading)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(header_text)
    document.add_paragraph("来源：模拟数据，仅作演示。", style="Source Note")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, required=True)
    tokens = load_word_tokens()
    parser.add_argument(
        "--font-profile",
        choices=tuple(tokens["font_profiles"]),
        default=tokens["default_font_profile"],
    )
    args = parser.parse_args()
    build(args.output, args.font_profile)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
