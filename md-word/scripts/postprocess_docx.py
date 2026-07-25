#!/usr/bin/env python3
"""Apply research-report paragraph, table, numbering, and TOC rules to a DOCX.

Pandoc is invoked without --toc/--number-sections; this step owns the Chinese
section numbering (第一部分：/1、) and inserts both directories (内容目录 +
图表目录) so heading text stays predictable for style matching.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

from build_reference_docx import configure_document, field, load_word_tokens


EXHIBIT_TITLE = re.compile(r"^图表\s*\d+\s*：")
SOURCE_NOTE = re.compile(r"^(?:资料)?来源\s*[：:]|^注\s*[：:]")
HEADING_STYLES = ("Heading 1", "Heading 2", "Heading 3")
UNNUMBERED_HEADINGS = {"摘要", "附录", "风险提示", "参考文献", "目录", "内容目录", "图表目录"}
CHINESE_DIGITS = "零一二三四五六七八九"


def chinese_numeral(number: int) -> str:
    if not 1 <= number <= 99:
        return str(number)
    tens, units = divmod(number, 10)
    if tens == 0:
        return CHINESE_DIGITS[units]
    text = ("" if tens == 1 else CHINESE_DIGITS[tens]) + "十"
    return text + (CHINESE_DIGITS[units] if units else "")


def _ensure(parent, tag: str):
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    return element


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = _ensure(tc_pr, "w:shd")
    shading.set(qn("w:fill"), fill)


def _page_break_before(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    _ensure(p_pr, "w:pageBreakBefore")


def _heading_text(paragraph) -> str:
    return paragraph.text.strip()


def _base_name(text: str) -> str:
    return text.split(" ", 1)[0].split("：", 1)[0].strip()


def _insert_paragraph_before(anchor, style_name: str, document) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    paragraph.style = document.styles[style_name]
    return paragraph


DROP_TITLE_SENTINEL = "__MD_WORD_DROP_TITLE__"


def _restyle_title(document: Document) -> None:
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() == DROP_TITLE_SENTINEL:
            paragraph._p.getparent().remove(paragraph._p)
    for paragraph in document.paragraphs:
        if paragraph.style.name in HEADING_STYLES:
            if paragraph.style.name == "Heading 1":
                paragraph.style = document.styles["Title"]
            return


def _insert_directories(document: Document) -> None:
    """Insert 内容目录 + 图表目录 after the 摘要 section (国金参考稿为双目录)."""
    if any(paragraph.text.strip() == "内容目录" for paragraph in document.paragraphs):
        return
    headings = [p for p in document.paragraphs if p.style.name in HEADING_STYLES]
    anchor = None
    for index, paragraph in enumerate(headings):
        if _base_name(_heading_text(paragraph)) == "摘要":
            anchor = headings[index + 1] if index + 1 < len(headings) else None
            break
    if anchor is None:
        anchor = headings[1] if len(headings) > 1 else None
    if anchor is None:
        return

    content_lines = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.style.name in HEADING_STYLES and paragraph.text.strip()
    ]
    exhibit_lines = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if EXHIBIT_TITLE.match(paragraph.text.strip())
    ]

    content_heading = _insert_paragraph_before(anchor, "Directory Heading", document)
    content_heading.add_run("内容目录")
    _page_break_before(content_heading)
    content_field = _insert_paragraph_before(anchor, "Normal", document)
    field(content_field, r'TOC \o "1-3" \h \z \u', content_lines or ["目录将在 Microsoft Word 中更新"])
    exhibit_heading = _insert_paragraph_before(anchor, "Directory Heading", document)
    exhibit_heading.add_run("图表目录")
    exhibit_field = _insert_paragraph_before(anchor, "Normal", document)
    field(
        exhibit_field,
        r'TOC \h \z \t "Exhibit Title,1"',
        exhibit_lines or ["图表目录将在 Microsoft Word 中更新"],
    )
    if anchor.style.name != "Heading 1":
        _page_break_before(anchor)


def _number_headings(document: Document) -> None:
    part = 0
    sub = 0
    for paragraph in document.paragraphs:
        if paragraph.style.name not in ("Heading 1", "Heading 2"):
            continue
        text = _heading_text(paragraph)
        if not text or _base_name(text) in UNNUMBERED_HEADINGS or not paragraph.runs:
            continue
        if re.match(r"^(?:第[一二三四五六七八九十]+部分：|\d+、)", text):
            continue
        first_run = paragraph.runs[0]
        if paragraph.style.name == "Heading 1":
            part += 1
            sub = 0
            first_run.text = f"第{chinese_numeral(part)}部分：{first_run.text}"
        else:
            sub += 1
            first_run.text = f"{sub}、{first_run.text}"


def postprocess(path: Path, font_profile: str = "preview") -> None:
    tokens = load_word_tokens()
    header_fill = tokens["colors"]["table_header"].lstrip("#")
    header_text = tokens["colors"]["table_header_text"].lstrip("#")
    document = Document(path)
    configure_document(document, font_profile)
    _restyle_title(document)
    _number_headings(document)
    _insert_directories(document)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if EXHIBIT_TITLE.match(text):
            paragraph.style = document.styles["Exhibit Title"]
        elif SOURCE_NOTE.match(text):
            paragraph.style = document.styles["Source Note"]
            paragraph.paragraph_format.keep_with_next = False
        elif _base_name(text) == "附录" and paragraph.style.name in HEADING_STYLES:
            paragraph.style = document.styles["Appendix Heading"]
        if paragraph._p.xpath(".//w:drawing") or paragraph._p.xpath(".//m:oMathPara"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            # 正文继承的 tokens 固定行距会把行内图片裁剪到一行高、把陈列
            # 公式的分数与上下限切掉半截，这类段落必须改回单倍自适应行距。
            paragraph.paragraph_format.line_spacing = tokens["image_line_spacing_multiple"]

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        if table.rows:
            header = table.rows[0]
            tr_pr = header._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
            for cell in header.cells:
                _shade(cell, header_fill)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(header_text)
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
        if table.rows:
            for cell in table.rows[-1].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True

    document.save(path)
