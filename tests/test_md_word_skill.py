import base64
import importlib.util
import json
import os
import shutil
import sys
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "md-word" / "scripts"
SHARED_SCRIPTS = ROOT / "shared-assets" / "scripts"
TOKENS = ROOT / "shared-assets" / "design-tokens.json"
REFERENCE = ROOT / "shared-assets" / "word" / "research-report-reference.docx"

# 1x1 transparent PNG so image fixtures are real, decodable files.
PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4nGNgYGBgAAAABQABh6FO1AAAAABJRU5ErkJggg=="
)

SAMPLE_MARKDOWN = r"""# 模拟策略研究

## 摘要

本文使用模拟数据。

## 方法

行内公式 $r_t=P_t/P_{t-1}-1$。

$$R=\sum_{t=1}^{T} r_t$$

## 回测

图表1：模拟组合净值

![模拟组合净值](chart.png)

来源：模拟数据，仅作演示。

图表2：模拟回测指标

| 指标 | 数值 |
|---|---:|
| 年化收益 | 10.2% |

来源：模拟数据，仅作演示。

## 结论

策略结果仅用于流程演示。

# 附录

参数口径。
"""


def load_module(name: str):
    path = SCRIPTS / f"{name}.py"
    if str(SCRIPTS) not in sys.path:
        sys.path.insert(0, str(SCRIPTS))
    if str(SHARED_SCRIPTS) not in sys.path:
        sys.path.insert(0, str(SHARED_SCRIPTS))
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def load_shared_module(name: str):
    path = SHARED_SCRIPTS / f"{name}.py"
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


class MdWordTokenTests(unittest.TestCase):
    def test_word_tokens_match_reference_scheme(self):
        tokens = json.loads(TOKENS.read_text(encoding="utf-8"))["word"]
        self.assertEqual(tokens["page"]["size"], "A4")
        self.assertEqual(tokens["margins_mm"], {"top": 31, "right": 12, "bottom": 16, "left": 12})
        self.assertEqual(tokens["font_sizes_pt"]["body"], 10.5)
        self.assertEqual(tokens["body_first_line_indent_chars"], 2)
        self.assertFalse(tokens["logo_slot"]["enabled"])

    def test_word_tokens_define_both_font_profiles(self):
        profiles = json.loads(TOKENS.read_text(encoding="utf-8"))["word"]["font_profiles"]
        self.assertEqual(set(profiles) >= {"preview", "delivery"}, True)
        self.assertEqual(profiles["delivery"]["body_zh"], "KaiTi")
        self.assertIn("Noto", profiles["preview"]["body_zh"])
        self.assertIn("fonts/NotoSansCJKsc-Regular.otf", profiles["preview"]["bundled_files"])
        self.assertIn("Microsoft YaHei", profiles["delivery"]["required_system_fonts"])

    def test_project_bundles_renderable_cjk_fonts(self):
        font_dir = ROOT / "shared-assets" / "fonts"
        self.assertGreater((font_dir / "NotoSansCJKsc-Regular.otf").stat().st_size, 10_000_000)
        self.assertGreater((font_dir / "NotoSerifCJKsc-Regular.otf").stat().st_size, 10_000_000)
        config = (font_dir / "fonts.conf").read_text(encoding="utf-8")
        self.assertIn("Noto Sans CJK SC", config)
        self.assertIn("__PROJECT_FONT_DIR__", config)
        self.assertIn("__FONT_CACHE_DIR__", config)
        self.assertNotIn("/Users/", config)
        self.assertNotIn("feishu", config)

    def test_word_font_preflight_rejects_missing_delivery_families(self):
        preflight = load_shared_module("font_preflight")
        def bundled_families(path: Path) -> set[str]:
            return {"Noto Serif CJK SC"} if "Serif" in path.name else {"Noto Sans CJK SC"}

        preview = preflight.check_font_profile(
            "word",
            "preview",
            system_families=set(),
            bundled_family_scanner=bundled_families,
        )
        delivery = preflight.check_font_profile(
            "word",
            "delivery",
            system_families={"Times New Roman", "Arial"},
            bundled_family_scanner=bundled_families,
        )
        self.assertEqual(preview["fonts"]["body_zh"], "Noto Serif CJK SC")
        self.assertTrue(preview["render_safe"])
        self.assertFalse(delivery["render_safe"])
        self.assertEqual(set(delivery["missing"]), {"KaiTi", "Microsoft YaHei"})

    def test_word_environment_report_includes_font_readiness(self):
        checker = load_module("check_environment")
        report = checker.build_report(
            "preview",
            tool_lookup=lambda name: f"/usr/bin/{name}",
            python_docx_available=True,
            system_families=set(),
            bundled_family_scanner=lambda path: {
                "Noto Serif CJK SC" if "Serif" in path.name else "Noto Sans CJK SC"
            },
        )
        self.assertTrue(report["ready"])
        self.assertTrue(report["visual_ready"])
        self.assertEqual(report["fonts"]["profile"], "preview")


class ValidateMarkdownTests(unittest.TestCase):
    def test_validator_accepts_unified_exhibit_numbering(self):
        validator = load_module("validate_markdown")
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "chart.png").write_bytes(PNG_BYTES)
            path = root / "report.md"
            path.write_text(SAMPLE_MARKDOWN, encoding="utf-8")
            result = validator.validate(path)
        self.assertEqual(result["errors"], [])
        self.assertEqual(result["math_expressions"], 2)
        self.assertEqual(result["exhibits"], 2)

    def test_validator_rejects_legacy_caption_style_and_broken_numbering(self):
        validator = load_module("validate_markdown")
        markdown = SAMPLE_MARKDOWN.replace("图表1：模拟组合净值", "图 1 模拟组合净值").replace(
            "图表2：模拟回测指标", "图表3：模拟回测指标"
        )
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "chart.png").write_bytes(PNG_BYTES)
            path = root / "report.md"
            path.write_text(markdown, encoding="utf-8")
            result = validator.validate(path)
        joined = "\n".join(result["errors"])
        self.assertIn("图表N：题", joined)
        self.assertIn("连续递增", joined)

    def test_validator_rejects_missing_source_and_unbalanced_math(self):
        validator = load_module("validate_markdown")
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "chart.png").write_bytes(PNG_BYTES)
            path = root / "report.md"
            path.write_text(
                "# 标题\n## 摘要\n$broken\n## 方法\n## 回测\n图表1：图题\n\n![图](chart.png)\n## 结论\n# 附录\n",
                encoding="utf-8",
            )
            result = validator.validate(path)
        joined = "\n".join(result["errors"])
        self.assertIn("公式分隔符不成对", joined)
        self.assertIn("缺少来源", joined)


class MaterializeReportTests(unittest.TestCase):
    def _payload(self) -> dict[str, object]:
        return {
            "title": "模拟策略研究",
            "data_status": "simulated",
            "summary": ["策略在模拟区间内录得正收益。"],
            "method": ["以期初净值归一为 1，按期比较策略与基准。"],
            "periods": ["2025Q1", "2025Q2", "2025Q3"],
            "series": [
                {"label": "策略净值", "values": [1.0, 1.08, 1.12]},
                {"label": "基准净值", "values": [1.0, 1.03, 1.05]},
            ],
            "metrics": [
                {"name": "累计收益", "value": "12.0%"},
                {"name": "最大回撤", "value": "3.1%"},
            ],
            "conclusion": ["结果只用于验证研究报告生成链路。"],
            "appendix": ["所有数值均为模拟口径。"],
            "sources": ["自建模拟数据集"],
        }

    def test_materializer_rejects_missing_required_fields(self):
        materializer = load_module("materialize_report")
        payload = self._payload()
        del payload["data_status"]
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "data.json"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "data_status"):
                materializer.materialize(source, Path(tmp) / "report.md")

    def test_materializer_is_deterministic_and_emits_valid_resources(self):
        materializer = load_module("materialize_report")
        validator = load_module("validate_markdown")
        payload = self._payload()
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "data.json"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            first = root / "first" / "report.md"
            second = root / "second" / "report.md"
            first_result = materializer.materialize(source, first)
            second_result = materializer.materialize(source, second)

            first_text = first.read_text(encoding="utf-8")
            second_text = second.read_text(encoding="utf-8")
            first_svg = (first.parent / "figures" / "series.svg").read_bytes()
            second_svg = (second.parent / "figures" / "series.svg").read_bytes()
            validation = validator.validate(first)

        self.assertEqual(first_text, second_text)
        self.assertEqual(first_svg, second_svg)
        self.assertIn(b'data-font-profile="preview"', first_svg)
        self.assertEqual(validation["errors"], [])
        self.assertEqual(validation["exhibits"], 2)
        self.assertIn("模拟数据，仅作演示", first_text)
        self.assertIn("![策略与基准序列](figures/series.svg)", first_text)
        self.assertEqual(first_result["exhibits"], 2)
        self.assertEqual(first_result["font_profile"], "preview")

    def test_materialized_svg_profile_must_match_docx_profile(self):
        materializer = load_module("materialize_report")
        builder = load_module("build_docx")
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "data.json"
            source.write_text(json.dumps(self._payload(), ensure_ascii=False), encoding="utf-8")
            markdown = root / "report.md"
            materializer.materialize(source, markdown, "preview")
            with self.assertRaisesRegex(RuntimeError, "SVG 字体档位与 DOCX 不一致"):
                builder.stage_source_with_png_figures(markdown, root / "stage", "delivery")


class ReferenceDocxTests(unittest.TestCase):
    def test_reference_docx_has_expected_page_and_blank_header(self):
        inspector = load_module("inspect_docx")
        doc = Document(REFERENCE)
        section = doc.sections[0]
        self.assertAlmostEqual(section.top_margin.mm, 31, places=1)
        self.assertAlmostEqual(section.left_margin.mm, 12, places=1)
        self.assertEqual("".join(p.text for p in section.header.paragraphs).strip(), "")
        style_names = [style.name for style in doc.styles]
        self.assertIn("Exhibit Title", style_names)
        self.assertIn("Directory Heading", style_names)
        with ZipFile(REFERENCE) as archive:
            footer_xml = "".join(
                archive.read(name).decode("utf-8")
                for name in archive.namelist()
                if name.startswith("word/footer") and name.endswith(".xml")
            )
        self.assertIn("PAGE", footer_xml)
        self.assertEqual(inspector.inspect(REFERENCE, font_profile="preview")["errors"], [])

    def test_body_styles_carry_cjk_first_line_indent(self):
        builder = load_module("build_reference_docx")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "ref.docx"
            builder.build(path)
            doc = Document(path)
            body_text = doc.styles["Body Text"]._element.xml
        self.assertIn('w:firstLineChars="200"', body_text)

    def test_delivery_profile_uses_windows_fonts(self):
        builder = load_module("build_reference_docx")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "ref.docx"
            builder.build(path, "delivery")
            doc = Document(path)
            normal_xml = doc.styles["Normal"]._element.xml
        self.assertIn("KaiTi", normal_xml)
        self.assertIn("Times New Roman", normal_xml)

    def test_reference_uses_token_driven_distances_and_style_spacing(self):
        builder = load_module("build_reference_docx")
        tokens = json.loads(TOKENS.read_text(encoding="utf-8"))
        tokens["word"]["header_distance_mm"] = 9.5
        tokens["word"]["footer_distance_mm"] = 7.25
        tokens["word"]["style_spacing_pt"] = {
            "Title": {"before": 0, "after": 21},
            "Heading 1": {"before": 19, "after": 9},
            "Heading 2": {"before": 13, "after": 7},
            "Heading 3": {"before": 9, "after": 5},
            "Exhibit Title": {"before": 6, "after": 3},
            "Source Note": {"before": 6, "after": 3},
            "Appendix Heading": {"before": 6, "after": 3},
            "Directory Heading": {"before": 6, "after": 3}
        }
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            custom_tokens = root / "tokens.json"
            custom_tokens.write_text(json.dumps(tokens, ensure_ascii=False), encoding="utf-8")
            original_tokens_path = builder.TOKENS_PATH
            builder.TOKENS_PATH = custom_tokens
            try:
                reference = root / "reference.docx"
                builder.build(reference)
            finally:
                builder.TOKENS_PATH = original_tokens_path
            document = Document(reference)

        section = document.sections[0]
        self.assertAlmostEqual(section.header_distance.mm, 9.5, places=1)
        self.assertAlmostEqual(section.footer_distance.mm, 7.25, places=1)
        self.assertAlmostEqual(document.styles["Title"].paragraph_format.space_after.pt, 21, places=1)
        self.assertAlmostEqual(document.styles["Heading 1"].paragraph_format.space_before.pt, 19, places=1)


class PostprocessTests(unittest.TestCase):
    def _pandoc_like_docx(self, path: Path) -> None:
        """Simulate the paragraph/style shape Pandoc emits before postprocess."""
        builder = load_module("build_reference_docx")
        document = Document()
        builder.configure_document(document)
        document.add_paragraph("模拟策略研究", style="Heading 1")
        document.add_paragraph("摘要", style="Heading 2")
        document.add_paragraph("本文使用模拟数据。", style="Body Text")
        document.add_paragraph("方法", style="Heading 2")
        document.add_paragraph("回测", style="Heading 2")
        document.add_paragraph("图表1：模拟回测指标", style="Body Text")
        table = document.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        document.add_paragraph("来源：模拟数据，仅作演示。", style="Body Text")
        document.add_paragraph("结论", style="Heading 2")
        document.add_paragraph("附录", style="Heading 1")
        document.save(path)

    def test_postprocess_relaxes_line_spacing_for_display_math(self):
        postprocessor = load_module("postprocess_docx")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "draft.docx"
            self._pandoc_like_docx(path)
            document = Document(path)
            math_paragraph = document.add_paragraph("", style="Body Text")
            math_para = OxmlElement("m:oMathPara")
            math = OxmlElement("m:oMath")
            math_para.append(math)
            math_paragraph._p.append(math_para)
            document.save(path)
            postprocessor.postprocess(path)
            doc = Document(path)
            spacing = [
                p.paragraph_format.line_spacing
                for p in doc.paragraphs
                if p._p.xpath(".//m:oMathPara")
            ]
        # 固定行距会把分数/上下限切掉半截，陈列公式段落必须回到单倍行距
        self.assertEqual(spacing, [1.0])

    def test_postprocess_numbers_headings_and_inserts_directories(self):
        postprocessor = load_module("postprocess_docx")
        inspector = load_module("inspect_docx")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "draft.docx"
            self._pandoc_like_docx(path)
            postprocessor.postprocess(path)
            doc = Document(path)
            texts = [p.text for p in doc.paragraphs]
            styles = {p.text: p.style.name for p in doc.paragraphs}
            document_xml = doc.element.xml
            inspection = inspector.inspect(path)

        self.assertEqual(styles["模拟策略研究"], "Title")
        self.assertIn("内容目录", texts)
        self.assertIn("图表目录", texts)
        self.assertLess(texts.index("内容目录"), texts.index("1、方法"))
        self.assertIn("1、方法", texts)
        self.assertIn("2、回测", texts)
        self.assertIn("3、结论", texts)
        self.assertEqual(styles["摘要"], "Heading 2")
        self.assertEqual(styles["附录"], "Appendix Heading")
        self.assertEqual(styles["图表1：模拟回测指标"], "Exhibit Title")
        self.assertEqual(styles["来源：模拟数据，仅作演示。"], "Source Note")
        self.assertIn(r'TOC \o "1-3"', document_xml)
        self.assertIn('TOC \\h \\z \\t "Exhibit Title,1"', document_xml)
        self.assertTrue(inspection["toc_cached"])
        self.assertTrue(inspection["exhibit_toc_cached"])

    def test_postprocess_is_idempotent(self):
        postprocessor = load_module("postprocess_docx")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "draft.docx"
            self._pandoc_like_docx(path)
            postprocessor.postprocess(path)
            postprocessor.postprocess(path)
            doc = Document(path)
            texts = [p.text for p in doc.paragraphs]
        self.assertEqual(texts.count("内容目录"), 1)
        self.assertIn("1、方法", texts)
        self.assertNotIn("1、1、方法", texts)


class InspectDocxTests(unittest.TestCase):
    def test_inspector_requires_native_omml_and_field_based_toc(self):
        inspector = load_module("inspect_docx")
        with tempfile.TemporaryDirectory() as tmp:
            fake = Path(tmp) / "fake.docx"
            with ZipFile(fake, "w") as archive:
                archive.writestr(
                    "word/document.xml",
                    '<w:document xmlns:w="w" xmlns:m="m"><w:body>'
                    '<w:fldChar w:fldCharType="begin"/><w:instrText>TOC \\o "1-3"</w:instrText>'
                    '<w:fldChar w:fldCharType="separate"/><w:t>摘要</w:t>'
                    '<w:fldChar w:fldCharType="end"/>'
                    "<m:oMath><m:r/></m:oMath></w:body></w:document>",
                )
                archive.writestr("word/footer1.xml", '<w:ftr xmlns:w="w"><w:instrText>PAGE</w:instrText></w:ftr>')
                archive.writestr(
                    "word/settings.xml",
                    '<w:settings xmlns:w="w"><w:updateFields w:val="true"/></w:settings>',
                )
            result = inspector.inspect(fake, expected_math=1)
        self.assertEqual(result["errors"], [])
        self.assertEqual(result["omml_count"], 1)

    def test_inspector_ignores_toc_as_plain_text_and_requires_exhibit_toc(self):
        inspector = load_module("inspect_docx")
        with tempfile.TemporaryDirectory() as tmp:
            fake = Path(tmp) / "fake.docx"
            with ZipFile(fake, "w") as archive:
                archive.writestr(
                    "word/document.xml",
                    '<w:document xmlns:w="w"><w:body><w:t>正文提到 TOC 与 PAGE 字样</w:t>'
                    "<w:t>图表1：示例</w:t></w:body></w:document>",
                )
                archive.writestr("word/footer1.xml", '<w:ftr xmlns:w="w"><w:t>PAGE</w:t></w:ftr>')
            result = inspector.inspect(fake)
        joined = "\n".join(result["errors"])
        self.assertIn("缺少 Word TOC 域", joined)
        self.assertIn("图表目录", joined)
        self.assertIn("页脚缺少 PAGE 域", joined)

    def test_inspector_rejects_visible_body_branding_and_empty_toc_cache(self):
        inspector = load_module("inspect_docx")
        with tempfile.TemporaryDirectory() as tmp:
            fake = Path(tmp) / "fake.docx"
            with ZipFile(fake, "w") as archive:
                archive.writestr(
                    "word/document.xml",
                    '<w:document xmlns:w="w"><w:body><w:p><w:r>'
                    '<w:fldChar w:fldCharType="begin"/><w:instrText>TOC \\o "1-3"</w:instrText>'
                    '<w:fldChar w:fldCharType="separate"/><w:fldChar w:fldCharType="end"/>'
                    '<w:t>国金证券</w:t></w:r></w:p></w:body></w:document>',
                )
                archive.writestr(
                    "word/footer1.xml",
                    '<w:ftr xmlns:w="w"><w:instrText>PAGE</w:instrText></w:ftr>',
                )
                archive.writestr(
                    "word/settings.xml",
                    '<w:settings xmlns:w="w"><w:updateFields w:val="true"/></w:settings>',
                )
            result = inspector.inspect(fake)
        joined = "\n".join(result["errors"])
        self.assertIn("内容目录域没有可见缓存结果", joined)
        self.assertIn("检测到禁止的可见品牌文字：国金证券", joined)


@unittest.skipUnless(
    shutil.which("pandoc") or shutil.which("quarto"),
    "需要 Pandoc（或 Quarto 内置 Pandoc）才能跑端到端转换",
)
class EndToEndBuildTests(unittest.TestCase):
    def test_build_docx_from_sample_markdown(self):
        builder = load_module("build_docx")
        renderer = load_module("render_docx_preview")
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "chart.png").write_bytes(PNG_BYTES)
            source = root / "report.md"
            source.write_text(SAMPLE_MARKDOWN, encoding="utf-8")
            output = root / "report.docx"
            result = builder.build(source, output)
            self.assertEqual(result["inspection"]["errors"], [])
            self.assertEqual(result["inspection"]["font_profile"], "preview")
            self.assertIn("Noto Sans CJK SC", result["inspection"]["expected_fonts"])
            rendered_pages = []
            if os.environ.get("RUN_OFFICE_E2E") == "1" and shutil.which("soffice") and shutil.which("pdftoppm"):
                rendered_pages = renderer.render(output, root / "pages", "preview")
            doc = Document(output)
            texts = [p.text for p in doc.paragraphs]
            titles = [p.text for p in doc.paragraphs if p.style.name == "Title"]
        self.assertIn("内容目录", texts)
        self.assertIn("图表目录", texts)
        self.assertTrue(any(t.startswith("1、") for t in texts))
        # HTML 中转的元数据标题必须被哨兵机制删除，只留报告标题本身
        self.assertEqual(titles, ["模拟策略研究"])
        self.assertFalse(any("DROP_TITLE" in t for t in texts))
        # implicit_figures 已禁用：alt 文本不得在图下再生成一份独立题注段落
        self.assertEqual(texts.count("模拟组合净值"), 0)
        self.assertIn("图表1：模拟组合净值", texts)
        if os.environ.get("RUN_OFFICE_E2E") == "1" and shutil.which("soffice") and shutil.which("pdftoppm"):
            self.assertGreater(len(rendered_pages), 0)


if __name__ == "__main__":
    unittest.main()
